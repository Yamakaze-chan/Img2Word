import cv2
import numpy as np
from PIL import Image
from PIL import ImageDraw, ImageFont, ImageFilter, ImageEnhance
import easyocr
import docx
from lib.vietocr.tool.predictor import Predictor
from lib.vietocr.tool.config import Cfg
from skimage.filters import threshold_yen
from preprocess.rotation_correction import rotate_image
from preprocess.brightness_and_contrast_correction import automatic_brightness_and_contrast

# -------------Preprocess-------------------
def bb_intersection_over_union(boxA, boxB):
    xA = max(boxA[0], boxB[0])
    yA = max(boxA[1], boxB[1])
    xB = min(boxA[2], boxB[2])
    yB = min(boxA[3], boxB[3])
    interArea = max(0, xB - xA +1) * max(0, yB - yA +1)
    boxAArea = (boxA[0] - boxA[2] -1) * (boxA[1] - boxA[3] -1)
    boxBArea = (boxB[0] - boxB[2] -1) * (boxB[1] - boxB[3] -1)
    iou = interArea / float(boxAArea + boxBArea - interArea)
    if ((boxA[0] - boxA[2] ) * (boxA[1] - boxA[3])>=(boxB[0] - boxB[2]) * (boxB[1] - boxB[3])):
        return iou, boxA[0], boxA[1], boxA[2], boxA[3]
    else:
        return iou, boxB[0], boxB[1], boxB[2], boxB[3]

def pre_process_image(img):
    blur = cv2.GaussianBlur(cv2.cvtColor(img, cv2.COLOR_BGR2GRAY), (5, 5), 0)
    _, thresh = cv2.threshold(blur, threshold_yen(img), 255, cv2.THRESH_BINARY)
    bitwise = cv2.bitwise_not(thresh)
    erosion = cv2.erode(bitwise, np.ones((1, 1) ,np.uint8), iterations=1)
    kernel_to_remove_gaps_between_words = np.array([
            [1,1,1,1,1,1,1,1],
            [1,1,1,1,1,1,1,1]
    ])
    dilation = cv2.dilate(erosion, np.ones((1, 1) ,np.uint8), iterations=1)
    return dilation

def pre_process_read_text(img):
    dilated_img = cv2.dilate(cv2.cvtColor(img, cv2.COLOR_BGR2GRAY), np.ones((5, 5), np.uint8))
    bg_img = cv2.medianBlur(dilated_img, 5)
    #--- finding absolute difference to preserve edges ---
    diff_img = 255 - cv2.absdiff(cv2.cvtColor(img, cv2.COLOR_BGR2GRAY), bg_img)
    #--- normalizing between 0 to 255 ---
    norm_img = cv2.normalize((diff_img), None, alpha=0, beta=255, norm_type=cv2.NORM_MINMAX, dtype=cv2.CV_8UC1)
    return cv2.threshold((norm_img), 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]

def sharpen(image):
    return np.array(ImageEnhance.Sharpness(Image.fromarray(image)).enhance(9.0))

def enhanced_edge(image):
    return np.array(Image.fromarray(image).filter(ImageFilter.EDGE_ENHANCE_MORE))

def grayscale(image):
    return cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
def threshold_image(image):
    return cv2.threshold(image, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
def yen_threshold_image(image):
    return cv2.threshold(image, threshold_yen(image), 255, cv2.THRESH_BINARY)[1]
def invert_image(image):
    return cv2.bitwise_not(image)
def dilation(image):
    kernel = np.ones((5, 5), np.uint8) 
    return cv2.dilate(image, kernel, iterations=1) 
def erosion(image):
    kernel = np.ones((5, 5), np.uint8) 
    return cv2.erode(image, kernel, iterations=1) 
def dilate_image(image):
    kernel_to_remove_gaps_between_words = np.array([
            [1,1,1],
            [1,1,1],
    ])
    dilated_image = cv2.dilate(image, kernel_to_remove_gaps_between_words, iterations=10)
    simple_kernel = np.ones((5,5), np.uint8)
    return cv2.dilate(dilated_image, simple_kernel, iterations=2)

def enhance_table_img(img):
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    bitwise = cv2.bitwise_not(gray)
    img = Image.fromarray(bitwise)
    enhancer = img.filter(ImageFilter.MaxFilter)
    # cv2.imshow("enhance table:", np.array(enhancer))
    # cv2.waitKey(0)
    return cv2.cvtColor(np.array(enhancer), cv2.COLOR_GRAY2BGR)

def preprocess_table_img(img):
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    thresh, img_bin = cv2.threshold(gray, threshold_yen(img), 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)
    # img_bin = 255-img_bin
    return img_bin

#-------------------Function-----------------------

def find_table(img):
    contours,_ = cv2.findContours(img,cv2.RETR_TREE, cv2.CHAIN_APPROX_SIMPLE)
    cordinates = []
    for cnt in contours:
        x,y,w,h = cv2.boundingRect(cnt)
        # cordinates.append((x,y,x+w,y+h))
        x_min, y_min, x_max, y_max = [x,y,x+w,y+h]
        # print(x, y , w, h)
        for temp_coor in cordinates:
            overlap = bb_intersection_over_union([x,y,x+w,y+h], temp_coor)
            if overlap[0] >= 0.0001 or (np.isclose(overlap[0], 0, rtol=1e-05, atol=1e-08, equal_nan=False) and overlap[0]!=0):
                x_min, y_min, x_max, y_max = overlap[1:]
                cordinates.remove(temp_coor)
        for temp_coor in cordinates:
            overlap = bb_intersection_over_union([x,y,x+w,y+h], temp_coor)
        cordinates.append([x_min, y_min, x_max, y_max])
        #bounding the images
    cv2.imwrite("result/tab.png", img[y_min:y_max, x_min: x_max])
    return cordinates

# def find_all_table(img):
#     coors = find_table(img)


def add_10_percent_padding(img):
    image_height = img.shape[0]
    padding = int(image_height * 0.1)
    return cv2.copyMakeBorder(img, padding, padding, padding, padding, cv2.BORDER_CONSTANT, value=[255, 255, 255])

def erode_vertical_lines(img):
    hor = np.array([[1,1,1,1,1,1]])
    vertical_lines_eroded_image = cv2.erode(invert_image(threshold_image((img))), hor, iterations=5)
    return cv2.dilate(vertical_lines_eroded_image, hor, iterations=10)

def erode_horizontal_lines(img):
    ver = np.array([[1],
            [1],
            [1],
            [1],
            [1],
            [1],
            [1]])
    horizontal_lines_eroded_image = cv2.erode(invert_image(threshold_image((img))), ver, iterations=5)
    return cv2.dilate(horizontal_lines_eroded_image, ver, iterations=10)

def combine_eroded_images(img):
    return cv2.add(erode_vertical_lines(img), erode_horizontal_lines(img))

def dilate_combined_image_to_make_lines_thicker(img):
    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 1))
    return cv2.dilate(combine_eroded_images(grayscale(img)), kernel, iterations=5)

def get_coor_vertical_lines(img):
    vertical_lines = []
    lines = cv2.HoughLinesP(erode_vertical_lines(grayscale(img)),rho=1,theta=np.pi/180,threshold=100,minLineLength=30,maxLineGap=3)
    for line in lines:
        for x1,y1,x2,y2 in line:
            for coor in vertical_lines:
                if(y1 == coor[1] and y2==coor[3]):
                    if (x2>=coor[2] and x1<=coor[0]):
                        vertical_lines.remove(coor)
                    elif (x2>=coor[2] and x1>coor[0] and x1<coor[2]):
                        x1 = coor[0]
                        vertical_lines.remove(coor)
                    elif (x2<coor[2] and x2>coor[0] and x1<=coor[0]):
                        x2 = coor[2]
                        vertical_lines.remove(coor)
            vertical_lines.append([int(x1),int(y1),int(x2),int(y2)])
            # cv2.line(img,(x1,y1),(x2,y2),(0,255,0),2) 
    # cv2.imshow("img", img) 
    # cv2.waitKey(0)
    return vertical_lines

def get_coor_horizontal_lines(img):
    horizontal_lines = []
    lines = cv2.HoughLinesP(erode_horizontal_lines(grayscale(img)),rho=1,theta=np.pi/180,threshold=100,minLineLength=30,maxLineGap=3)
    for line in lines:
        for x2,y2,x1,y1 in line:
            for coor in horizontal_lines:
                if(x1 == coor[0] and x2==coor[2]):
                    if (y2>=coor[3] and y1<=coor[1]):
                        horizontal_lines.remove(coor)
                    elif (y2>=coor[3] and y1>coor[1] and y1<coor[3]):
                        y1 = coor[1]
                        horizontal_lines.remove(coor)
                    elif (y2<coor[3] and y2>coor[1] and y1<=coor[1]):
                        y2 = coor[3]
                        horizontal_lines.remove(coor)
            horizontal_lines.append([int(x1),int(y1),int(x2),int(y2)])
            # cv2.line(img,(x1,y1),(x2,y2),(0,255,0),2)
    # cv2.imshow("img", img) 
    # cv2.waitKey(0)
    return horizontal_lines

def get_horizontal_lines(img_bin):
    kernel_len = img_bin.shape[1]//120
    if(kernel_len == 0):
        return [None, None]
    hor_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (kernel_len, 1))
    image_horizontal = cv2.erode(img_bin, hor_kernel, iterations=3)
    horizontal_lines = cv2.dilate(image_horizontal, hor_kernel, iterations=3)
    h_lines = cv2.HoughLinesP(horizontal_lines, 1,np.pi/180,100,0,10)
    return [h_lines, kernel_len]

def get_vertical_lines(img_bin):
    kernel_len = img_bin.shape[1]//120
    if(kernel_len == 0):
        return [None, None]
    ver_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, kernel_len))
    image_vertical = cv2.erode(img_bin, ver_kernel, iterations=3)
    vertical_lines = cv2.dilate(image_vertical, ver_kernel, iterations=3)
    v_lines = cv2.HoughLinesP(vertical_lines, 1,np.pi/180,100,0,10)
    return [v_lines, kernel_len]
    
def group_h_lines(val):
    h_lines, thin_thresh = val
    new_h_lines = []
    if h_lines is not None:
        while len(h_lines) > 0:
            thresh = sorted(h_lines, key=lambda x: x[0][1])[0][0]
            lines = [line for line in h_lines if thresh[1] -
                    thin_thresh <= line[0][1] <= thresh[1] + thin_thresh]
            h_lines = [line for line in h_lines if thresh[1] - thin_thresh >
                    line[0][1] or line[0][1] > thresh[1] + thin_thresh]
            x = []
            for line in lines:
                x.append(line[0][0])
                x.append(line[0][2])
            x_min, x_max = min(x) - int(5*thin_thresh), max(x) + int(5*thin_thresh)
            new_h_lines.append([x_min, thresh[1], x_max, thresh[1]])
        return new_h_lines
    else:
        return None

def group_v_lines(val):
    v_lines, thin_thresh = val
    new_v_lines = []
    if v_lines is not None:
        while len(v_lines) > 0:
            thresh = sorted(v_lines, key=lambda x: x[0][0])[0][0]
            lines = [line for line in v_lines if thresh[0] -
                    thin_thresh <= line[0][0] <= thresh[0] + thin_thresh]
            v_lines = [line for line in v_lines if thresh[0] - thin_thresh >
                    line[0][0] or line[0][0] > thresh[0] + thin_thresh]
            y = []
            for line in lines:
                y.append(line[0][1])
                y.append(line[0][3])
            y_min, y_max = min(y) - int(4*thin_thresh), max(y) + int(4*thin_thresh)
            new_v_lines.append([thresh[0], y_min, thresh[0], y_max])
        return new_v_lines
    else:
        return None

def line_intersection(line1, line2):
    xdiff = (line1[0][0] - line1[1][0], line2[0][0] - line2[1][0])
    ydiff = (line1[0][1] - line1[1][1], line2[0][1] - line2[1][1])

    def det(a, b):
        return a[0] * b[1] - a[1] * b[0]

    div = det(xdiff, ydiff)
    if div == 0:
       raise Exception('lines do not intersect')

    d = (det(*line1), det(*line2))
    x = det(d, xdiff) / div
    y = det(d, ydiff) / div
    return x, y

def get_bottom_right(right_points, bottom_points, points):
    for right in right_points:
        for bottom in bottom_points:
            if [right[0], bottom[1]] in points:
                return right[0], bottom[1]
    return None, None

def get_cell_span(cell_coor, list_of_h_lines, list_of_v_lines):
    sorted_h_lines = sorted(list_of_h_lines, key=lambda x: x[0], reverse=False)
    sorted_v_lines = sorted(list_of_v_lines, key=lambda x: x[1], reverse=False)
    new_coor = []
    for coor in cell_coor:
        colspan = 1
        rowspan = 1
        for i_h in range(0,len(sorted_h_lines)):
            if coor[0] == sorted_h_lines[i_h][0]:
                for temp_i_h in range(i_h+1, len(sorted_h_lines)):
                    if coor[2] == sorted_h_lines[temp_i_h][2]:
                        break
                    colspan = colspan + 1
                break
        for i_v in range(0,len(sorted_v_lines)):
            if coor[1] == sorted_v_lines[i_v][1]:
                
                for temp_i_v in range(i_v+1, len(sorted_v_lines)):
                    if coor[3] == sorted_v_lines[temp_i_v][3]:
                        break
                    rowspan = rowspan + 1
                break
        new_coor.append([coor[0], coor[1], coor[2], coor[3], colspan, rowspan])
    return new_coor

def get_list_of_cell_table(img, list_of_h_lines, list_of_v_lines, width, heigth, percent=1.0):
    points = []
    # image_temp = img.copy()
    for i in list_of_h_lines:
        for j in list_of_h_lines:
            if i!=j and (abs(i[1] - j[1]) < percent*width/100):
                list_of_h_lines.remove(j)
    for k in list_of_v_lines:
        for m in list_of_v_lines:
            if k!=m and (abs(k[0] - m[0]) < percent*heigth/100):
                list_of_v_lines.remove(m)
    for hline in list_of_h_lines:
        for vline in list_of_v_lines:
            line1 = [np.array([hline[0], hline[1]]), np.array([hline[2], hline[3]])]
            line2 = [np.array([vline[0], vline[1]]), np.array([vline[2], vline[3]])]
            x , y = line_intersection(line1, line2)
            if (hline[0]==hline[2]==x and vline[1] == vline[3]==y and vline[0]<=x<=vline[2] and hline[1]<=y<=hline[3]) or (hline[1]==hline[3]==y and vline[0]==vline[2]==x and vline[1]<=y<=vline[3] and hline[0]<=x<=hline[2]):
                points.append([int(x), int(y)])
        #         print("line1:")
        #         print(line1)
        #         print("line2:")
        #         print(line2)
        #         print("x: " + str(x) + " y: " + str(y))
        #         cv2.circle(image_temp, (int(x), int(y)), radius=2, color=(0, 0, 255), thickness=-5)
        # cv2.imshow("img", image_temp) 
        # cv2.waitKey(0)
    cells = []
    for point in points:
        left, top = point
        right_points = sorted(
            [p for p in points if p[0] > left and p[1] == top], key=lambda x: x[0])
        bottom_points = sorted(
            [p for p in points if p[1] > top and p[0] == left], key=lambda x: x[1])
        right, bottom = get_bottom_right(
            right_points, bottom_points, points)
        if right and bottom:
            # cv2.rectangle(img, (left, top), (right, bottom), (0, 0, 255), 1)
            # print("ngang:" + str((right-left)*100/width))
            # print("dọc:" + str((bottom - top)*100/heigth))
            # print(left, top, right, bottom)
            cells.append([left, top, right, bottom])
    
    return cells

def text_to_para(list_of_text):
    list_of_text = sorted(list_of_text, key=lambda x: x[0][0][1], reverse=False)
    middle_line = int(list_of_text[0][0][2][1] + list_of_text[0][0][0][1])/2
    string_of_text = []
    line = []
    for text_string in list_of_text:
        if(text_string[0][0][1] < middle_line):
            line.append(text_string)
        else:
            line = sorted(line, key=lambda x: x[0][0][0], reverse=False)
            middle_line = int(text_string[0][2][1] + text_string[0][0][1])/2
            string_of_text.append(' '.join(x[1] for x in line))
            line = []
            line.append(text_string)
        if text_string == list_of_text[-1]:
            line = sorted(line, key=lambda x: x[0][0][0], reverse=False)
            string_of_text.append(' '.join(x[1] for x in line))
    return '\n'.join(string_of_text)

def rearrange_cell(list_of_cell):
    return_sorted_cell = []
    vertical_sort = sorted(list_of_cell, key=lambda x: x[1], reverse=False)
    horizontal_sort = []
    y_line = list_of_cell[0][1]
    for cell in vertical_sort:
        if(cell[1] == y_line):
            horizontal_sort.append(cell)
        else:
            horizontal_sort = sorted(horizontal_sort, key=lambda x: x[0], reverse=False)
            y_line = cell[1]
            return_sorted_cell = return_sorted_cell + horizontal_sort
            horizontal_sort = []
            horizontal_sort.append(cell)
        if cell == vertical_sort[-1]:
            horizontal_sort = sorted(horizontal_sort, key=lambda x: x[0], reverse=False)
            return_sorted_cell = return_sorted_cell + horizontal_sort
    return return_sorted_cell

def get_cell_bounding_box(img, width, heigth, percent=1.0):
    temp_img = preprocess_table_img(enhance_table_img(img))
    h_lines = group_h_lines(get_horizontal_lines(temp_img))
    v_lines = group_v_lines(get_vertical_lines(temp_img))
    # print(h_lines, v_lines)
    if h_lines is None or v_lines is None:
        return [img, None, None, 0, 0]
    all_cell = get_list_of_cell_table(img, h_lines, v_lines, width, heigth, percent)
    if all_cell:
        cell_data = list(zip(*all_cell))
        x_min_coor = np.amin(cell_data[0], axis=0)
        y_min_coor = np.amin(cell_data[1], axis=0)
        x_max_coor = np.amax(cell_data[2], axis=0)
        y_max_coor = np.amax(cell_data[3], axis=0)
        list_of_cell = []
        try:
            add_border_img = cv2.copyMakeBorder(img[y_min_coor:y_max_coor, x_min_coor:x_max_coor], 0, 1, 0, 1, cv2.BORDER_CONSTANT, value=[0, 0, 0])
            coor_vlines = get_coor_vertical_lines(add_border_img)
            coor_hlines = get_coor_horizontal_lines(add_border_img)
            list_of_cell = get_list_of_cell_table(add_border_img, coor_vlines, coor_hlines, width, heigth, percent)
            # for cell in list_of_cell:
            #     cv2.rectangle(add_border_img,(cell[0],cell[1]),(cell[2],cell[3]),(0,0,255),1)
            # cv2.imshow("img", add_border_img) 
            # cv2.waitKey(0) 
            list_of_cell = rearrange_cell(list_of_cell)
            list_of_cell = get_cell_span(list_of_cell,coor_hlines,coor_vlines)
            col_num = len(coor_vlines)-1
            row_num = len(coor_hlines)-1
        except Exception as e:
            print(e)
            print("Not table")
            return [img, None, [], 0, 0]
        return [img, [x_min_coor, y_min_coor, x_max_coor, y_max_coor], list_of_cell, col_num, row_num]
    else:
        return [img, None, [], 0 , 0]
    
def get_alignment(left, right, width):
    align = 'left'
    percent = 20
    if 0 <= left < percent*width/100:
        if (100-percent)*width/100 < right <= width:
            align = 'left'
        else:
            align = 'left'
    else:
        if (100-percent)*width/100 < right <= width:
            align = 'right'
        else:
            align = 'center'
    return align

def read_text_of_cell(model, read_text_model, img, data_of_cell, width = -1, percent = 0.0):
    # cv2.imshow("bbox_text", img)
    # cv2.waitKey(0)
    image_to_read = cv2.cvtColor(yen_threshold_image(cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)), cv2.COLOR_GRAY2BGR)
    if data_of_cell is not None:
        for data in data_of_cell:
            # print("************")
            list_of_text = []
            for text_bbox in model.detect((img[data[1]:data[3],data[0]:data[2]])):
                # print(text)
                x = int(data[0]+text_bbox[0][0][0])
                y = int(data[1]+text_bbox[0][0][2])
                x_max = int(data[0]+text_bbox[0][0][1])
                y_max = int(data[1]+text_bbox[0][0][3])
                if (y_max - y > 0 and x_max - x > 0):
                    text = read_text_model.predict(Image.fromarray(rotate_image(image_to_read[y:y_max, x:x_max])[0]), return_prob=False)
                # cv2.rectangle(img,(x,y),(x_max,y_max),(0,0,255),1)
                # print(text)
                    list_of_text.append(text)
                # cv2.imshow("img", img) 
                # cv2.waitKey(0) 
            if len(list_of_text) !=0:
                # print(text_to_para(list_of_text))
                return ' '.join(list_of_text)
            # if len(list_of_text) > 1:
            # print(data[4])
            # print(data[5])
            # cv2.rectangle(img,(data[0],data[1]),(data[2],data[3]),(0,0,255),1)
            # cv2.imshow("img", img) 
            # cv2.waitKey(0) 
    else:
        list_of_text = []
        left_text = img.shape[1]
        right_text = 0
        # for text_bbox in model.detect((img),height_ths=0.8,width_ths=0.9,text_threshold=0.1):
        #     if len(text_bbox[0]) > 0:
        #         print(text_bbox)
        #         if isinstance(text_bbox[0][0][0], list):
        #             x = int(min(text_bbox[0][0][0][0],text_bbox[0][0][3][0])) if int(min(text_bbox[0][0][0][0],text_bbox[0][0][3][0])) > 0 else 0
        #             y = int(min(text_bbox[0][0][0][1],text_bbox[0][0][1][1])) if int(min(text_bbox[0][0][0][1],text_bbox[0][0][1][1])) > 0 else 0
        #             x_max = int(max(text_bbox[0][0][1][0],text_bbox[0][0][2][0])) if int(max(text_bbox[0][0][1][0],text_bbox[0][0][2][0])) < img.shape[1] else img.shape[1]-1
        #             y_max = int(max(text_bbox[0][0][2][1],text_bbox[0][0][3][1])) if int(max(text_bbox[0][0][2][1],text_bbox[0][0][3][1])) < img.shape[0] else img.shape[0]-1
        #         else:
        #             x = int(text_bbox[0][0][0]) if int(text_bbox[0][0][0]) > 0 else 0
        #             y = int(text_bbox[0][0][2]) if int(text_bbox[0][0][2]) > 0 else 0
        #             x_max = int(text_bbox[0][0][1]) if int(text_bbox[0][0][1]) < img.shape[1] else img.shape[1]-1
        #             y_max = int(text_bbox[0][0][3]) if int(text_bbox[0][0][2]) < img.shape[0] else img.shape[0]-1
        #         # cv2.imshow("img", img[y:y_max, x:x_max]) 
        #         # cv2.waitKey(0) 
        #         text = read_text_model.predict(Image.fromarray((image_to_read[y:y_max, x:x_max])), return_prob=False)
        #         if x < left_text:
        #             left_text = x
        #         if right_text < x_max:
        #             right_text = x_max
        #         # cv2.rectangle(img,(x,y),(x_max,y_max),(0,0,255),1)
        #         # print(text)
        #         list_of_text.append(text)
        #         # cv2.imshow("img", img) 
        #         # cv2.waitKey(0) 
        # if len(list_of_text) !=0:
        #     # print(text_to_para(list_of_text))
        #     return [' '.join(list_of_text), get_alignment(left_text, right_text, width)]
        text = []
        text_image = rotate_image(img)[0]
        for text_coors in get_line_of_text_from_coor(text_image, text_extraction(model, text_image, [])[0]):
            for text_coor in text_coors[1]:
                # print(text_coor)
                if(text_coor[3] - text_coor[1] > 0 and text_coor[2] - text_coor[0] > 0):
                    # print(text_coor[3], text_coor[1], text_coor[2] , text_coor[0])
                    text.append(read_text_model.predict(Image.fromarray(text_image[text_coor[1]:text_coor[3], text_coor[0]:text_coor[2]]), return_prob=False))
        # print(text)
        # cv2.imshow("img", img) 
        # cv2.waitKey(0) 
        return ' '.join(text)


def ResizeWithAspectRatio(image, width=None, height=None, inter=cv2.INTER_AREA):
    dim = None
    (h, w) = image.shape[:2]
    if width is None and height is None:
        return image
    if width is None:
        r = height / float(h)
        dim = (int(w * r), height)
    else:
        r = width / float(w)
        dim = (width, int(h * r))
    return cv2.resize(image, dim, interpolation=inter)


def find_contours(image):
    return cv2.findContours(pre_process_image(image), cv2.RETR_TREE, cv2.CHAIN_APPROX_SIMPLE)[0]

def convert_contours_to_bounding_boxes(model, image):
    bounding_boxes = []
    # for contour in find_contours(image):
    #     x, y, w, h = cv2.boundingRect(contour)
    #     if w < image.shape[1]*1.0/100 or h < image.shape[0]*1.0/100:
    #         continue
    #     x = x - int(1.0*image.shape[1]/100) if x > int(1.0*image.shape[1]/100) else 0
    #     y = y - int(1.0*image.shape[0]/100) if y > int(1.0*image.shape[0]/100) else 0
    #     x_max = x+w
    #     x_max = x_max + int(1.0*image.shape[1]/100) if x_max < (image.shape[1] - int(1.0*image.shape[1]/100)) else image.shape[1]-1
    #     y_max = y+h
    #     y_max = y_max + int(1.0*image.shape[0]/100) if y_max < (image.shape[0] - int(1.0*image.shape[0]/100)) else image.shape[0]-1
    #     bounding_boxes.append((x, y, x_max, y_max))
    for contour in model.readtext(sharpen(automatic_brightness_and_contrast(image)[0])):
        x = contour[0][0][0]
        y = contour[0][0][1]
        # if w < image.shape[1]*1.0/100 or h < image.shape[0]*1.0/100:
        #     continue
        x = x - int(1.0*image.shape[1]/100) if x > int(1.0*image.shape[1]/100) else 0
        # y = y - int(1.0*image.shape[0]/100) if y > int(1.0*image.shape[0]/100) else 0
        x_max = contour[0][2][0]
        x_max = x_max + int(1.0*image.shape[1]/100) if x_max < (image.shape[1] - int(1.0*image.shape[1]/100)) else image.shape[1]-1
        y_max = contour[0][2][1]
        # y_max = y_max + int(1.0*image.shape[0]/100) if y_max < (image.shape[0] - int(1.0*image.shape[0]/100)) else image.shape[0]-1
        bounding_boxes.append((int(x), int(y), int(x_max), int(y_max)))
    for index_1, bbox in enumerate(bounding_boxes):
        for index_2, temp_bbox in enumerate(bounding_boxes):
            #print(bb_intersection_over_union(bbox, temp_bbox)[0])
            iou = bb_intersection_over_union(bbox, temp_bbox)
            if iou[0]>0 and iou[0]!=1.0 and ((bbox[1] < ((temp_bbox[3]+temp_bbox[1])/2) < bbox[3]) or (temp_bbox[3] > ((bbox[1] + bbox[3])/2) > temp_bbox[1])):
                if ((bbox[0] - bbox[2] ) * (bbox[1] - bbox[3])>=(temp_bbox[0] - temp_bbox[2]) * (temp_bbox[1] - temp_bbox[3])):
                    bounding_boxes[index_2] = bbox
                else:
                    bounding_boxes[index_1] = bbox
    bounding_boxes=list(set(tuple(element) for element in bounding_boxes))
    # for x, y, x_max, y_max in bounding_boxes:
    #     cv2.rectangle(image, (x, y), (x_max, y_max), (0, 255, 0), 1)
    return bounding_boxes

def rearrange_list_of_text(list_of_text_bbox):
    return sorted(sorted(list_of_text_bbox, key=lambda x: x[0], reverse=False), key=lambda x: x[1], reverse=False)

def get_margin_of_paper(list_of_text_bbox):
    left = min(x[0] for x in list_of_text_bbox)
    top = min(y[1] for y in list_of_text_bbox)
    right = max(x_max[2] for x_max in list_of_text_bbox)
    bottom = max(y_max[3] for y_max in list_of_text_bbox)
    return [left, top, right, bottom]

def add_data_for_text_coor(list_of_text_bbox):
    middle_line = int(list_of_text_bbox[0][1] + list_of_text_bbox[0][3])/2
    list_of_lines = []
    line = []
    # print(list_of_text_bbox)
    for text_string in list_of_text_bbox:
        if(text_string[1] < middle_line):
            line.append(text_string)
        else:
            line = sorted(line, key=lambda x: x[0], reverse=False)
            middle_line = int(text_string[1] + text_string[3])/2
            list_of_lines.append(line)
            line = []
            line.append(text_string)
        if text_string == list_of_text_bbox[-1]:
            line = sorted(line, key=lambda x: x[0], reverse=False)
            list_of_lines.append(line)
    return list_of_lines

def merge_bbox_on_the_same_line(img, list_of_bbox):
    dilated_img = pre_process_image(img)
    # cv2.imshow('merge bbox', np.concatenate((img, cv2.cvtColor(dilated_img, cv2.COLOR_GRAY2BGR)), axis=1 ))
    # cv2.waitKey(0)
    for bbox in list_of_bbox:
        middle_line = int((bbox[3]-bbox[1])/2)
        # print("___________")
        # print(bbox)
        # print(dilated_img.shape)
        while bbox[0] > 0 and dilated_img[middle_line][bbox[0]] >100:
            bbox[0] = bbox[0]-1
        while bbox[2] < img.shape[1]-1 and dilated_img[middle_line][bbox[2]] > 100:
            bbox[2] = bbox[2]+1
    for index_1, bbox in enumerate(list_of_bbox):
        for index_2, temp_bbox in enumerate(list_of_bbox):
            #print(bb_intersection_over_union(bbox, temp_bbox)[0])
            iou = bb_intersection_over_union(bbox, temp_bbox)
            if iou[0]>0 and iou[0]!=1.0:
                list_of_bbox[index_1] = [min(bbox[0], temp_bbox[0]), min(bbox[1], temp_bbox[1]), max(bbox[2], temp_bbox[2]),max(bbox[3], temp_bbox[3])]
                list_of_bbox[index_2] = [min(bbox[0], temp_bbox[0]), min(bbox[1], temp_bbox[1]), max(bbox[2], temp_bbox[2]),max(bbox[3], temp_bbox[3])]
    list_of_bbox=list(set(tuple(element) for element in list_of_bbox))
    list_of_bbox=sorted(list_of_bbox, key=lambda x: x[0], reverse=False)
    return list_of_bbox

def table_extraction(img, list_all_coor, percent = 1.0):
    image = pre_process_image(img)
    cv2.imwrite('result/table.png', image) 
    image_without_table = img.copy()
    table_coor = None
    for j in find_table(image):
        if(j[2] - j[0])*(j[3] - j[1]) > (percent*image.shape[1]*image.shape[0]/100):
                # cv2.imshow("table", img[j[1]:j[3]+3, j[0]:j[2]+3])
                # cv2.waitKey(0)
                [cell_img, table_coor, list_cell_coor, col, row] = get_cell_bounding_box(img[j[1]:j[3], j[0]:j[2]], image.shape[1], image.shape[0], percent)
                if table_coor is not None:
                    list_all_coor.append([(table_coor[0] + j[0], table_coor[1] + j[1], table_coor[2] + j[0], table_coor[3] + j[1]), list_cell_coor, col, row])
                    # read_text_of_cell(model, cell_img[table_coor[1]:table_coor[3],table_coor[0]:table_coor[2]], list_cell_coor)
                    cv2.rectangle(image_without_table, (j[0]+table_coor[0], j[1]+table_coor[1]), (j[0]+table_coor[2]+3, j[1]+table_coor[3]+3), (255, 255, 255), -1)
    return [image_without_table, list_all_coor]

def text_extraction(model, img, list_all_coor):
    margin_of_paper = [0, 0 , img.shape[1], img.shape[0]]
    text_coordinates = rearrange_list_of_text(convert_contours_to_bounding_boxes(model, img))
    if len(text_coordinates)> 0:
        margin_of_paper = get_margin_of_paper(text_coordinates)
        for temp in add_data_for_text_coor(text_coordinates):
            # print(temp)
            if(len(temp) > 1):
                x = min(x[0] for x in temp)
                y = min(y[1] for y in temp)
                x_max = max(x_max[2] for x_max in temp)
                y_max = max(y_max[3] for y_max in temp)
                temp = [list(temp_) for temp_ in temp]
                for temp_x, temp_y, temp_x_max, temp_y_max in temp:
                    if (temp_x_max - temp_x < img.shape[1]*0.01) and (temp_y_max - temp_y < img.shape[0]*0.01):
                        temp.remove([temp_x, temp_y, temp_x_max, temp_y_max])
                list_all_coor.append([(x, y, x_max, y_max), temp, None, None])
                # cv2.rectangle(image_without_table, (x,y), (x_max,y_max), (255, 0, 0) , 1)
            else:
                if (temp[0][2] - temp[0][0] > img.shape[1]*0.01) and (temp[0][3]- temp[0][1] > img.shape[0]*0.01):
                    list_all_coor.append([temp[0], [[temp[0][0], temp[0][1],temp[0][2],temp[0][3]]], None, None])
                # cv2.rectangle(image_without_table, (temp[0][0],temp[0][1]), (temp[0][2],temp[0][3]), (255, 0, 0) , 1)
    # print("_____________________________")
    # for i in list_all_coor:
    #     print(i)
    return [list_all_coor, margin_of_paper]

def get_line_of_text_from_coor(img, list_all_coor):
    for index in range(0,len(list_all_coor)):
        if(list_all_coor[index][-1] is None and list_all_coor[index][-2] is None):
            if list_all_coor[index][0][1] > 0:
            # print(list_all_coor[index])
                list_all_coor[index][1] = merge_bbox_on_the_same_line(img[list_all_coor[index][0][1]:list_all_coor[index][0][3], 0:img.shape[1]-1], list_all_coor[index][1])
            else:
                list_all_coor[index][1] = merge_bbox_on_the_same_line(img[0:list_all_coor[index][0][3], 0:img.shape[1]-1], list_all_coor[index][1])
    return list_all_coor  

def make_doc_file(img, paper_list_coor,detect_text_model, read_text_model, doc, margin_of_paper):
    if len(paper_list_coor)==0:
        # print("Cannot convert to file")
        raise Exception("Không thể chuyển thành file Word")
    else:
        paper_list_coor = sorted(paper_list_coor, key=lambda x: x[0][1], reverse=False)
        i = 0
        for coor in paper_list_coor:
            print("Progress: " +str(i)+ "/"+ str(len(paper_list_coor)))
            i+=1
            # ADD TABLE
            if(coor[-1] is not None and coor[-2] is not None):
                print(coor)
                i=0
                table = doc.add_table(rows=coor[-2], cols=coor[-1], style="Table Grid") 
                print("Row: " + str(coor[-2]) + ",Col: " + str(coor[-1]))
                for row_index in range(0,coor[-2]):
                    col_index = 0
                    while col_index <coor[-1]:
                        if row_index > 0:
                            if table.cell(row_index, col_index)._tc == table.cell(row_index-1, col_index)._tc:
                                col_index += 1
                                continue
                        if col_index > 0:
                            if table.cell(row_index, col_index)._tc == table.cell(row_index, col_index-1)._tc:
                                col_index +=1
                                continue
                        print(coor[1][i])
                        if coor[1][i][-2] != 1 and coor[1][i][-1] != 1:
                            table.cell(row_index, col_index).merge(table.cell(row_index+coor[1][i][-1]-1, col_index+coor[1][i][-2]-1))
                            col_index = col_index+coor[1][i][-2]-1
                        elif coor[1][i][-1] != 1 and coor[1][i][-2] == 1:
                            table.cell(row_index, col_index).merge(table.cell(row_index+coor[1][i][-1]-1, col_index))
                        elif coor[1][i][-1] == 1 and coor[1][i][-2] != 1:
                            table.cell(row_index, col_index).merge(table.cell(row_index, col_index+coor[1][i][-2]-1))
                            col_index = col_index+coor[1][i][-2]-1
                        col_index += 1
                        i = i + 1
                i=0
                for row_index in range(0,coor[-2]):
                    col_index = 0
                    while col_index < coor[-1]:
                        if row_index > 0:
                            if table.cell(row_index, col_index)._tc == table.cell(row_index-1, col_index)._tc:
                                col_index += 1
                                continue
                        if col_index > 0:
                            if table.cell(row_index, col_index)._tc == table.cell(row_index, col_index-1)._tc:
                                col_index +=1
                                continue
                        row = table.rows[row_index]
                        cell_alignment = get_alignment(coor[0][0] + coor[1][i][0]+1,coor[0][0] + coor[1][i][2]-1, coor[1][i][2]-coor[1][i][0])
                        row.cells[col_index].text = read_text_of_cell(detect_text_model,read_text_model, img[coor[0][1] + coor[1][i][1]:coor[0][1] + coor[1][i][3],coor[0][0] + coor[1][i][0]:coor[0][0] + coor[1][i][2]], None, coor[1][i][2]-coor[1][i][0]) or ""
                        p = row.cells[col_index].paragraphs[0]
                        if cell_alignment == 'left':
                            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
                        elif cell_alignment == 'right':
                            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                        else:
                            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                        col_index += 1
                        i = i + 1
            else:
                #ADD TEXT
                if len(coor[1]) == 1:
                    alignment = get_alignment(coor[1][0][0],coor[1][0][2], img.shape[1])
                    doc_para = doc.add_paragraph(read_text_of_cell(detect_text_model,read_text_model, img[coor[1][0][1]:coor[1][0][3],coor[1][0][0]:coor[1][0][2]], None, margin_of_paper[2]-margin_of_paper[0])  or "")
                    if alignment == 'left':
                        doc_para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
                    elif alignment == 'right':
                        doc_para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        doc_para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                else:
                    text_on_same_line = doc.add_table(rows=1, cols=len(coor[1]))
                    for index in range(0, len(coor[1])):
                        cell = text_on_same_line.rows[0].cells
                        alignment = get_alignment(int(coor[1][index][0]),int(coor[1][index][2]), img.shape[1])
                        cell[index].text = read_text_of_cell(detect_text_model,read_text_model, img[int(coor[1][index][1]):int(coor[1][index][3]),int(coor[1][index][0]):int(coor[1][index][2])], None, coor[1][index][2]-coor[1][index][0]) or ""
                        p = cell[index].paragraphs[0]
                        if alignment == 'left':
                            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
                        elif alignment == 'right':
                            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
                        else:
                            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    return doc

def from_image_to_docx(detect_text_model,read_text_model, img,file_name,save_dir):
    image = pre_process_image(img)
    image_without_table = img.copy()
    image_without_line_of_text = img.copy()
    table_coor = text_coordinates = None
    paper_list_coor = []
    margin_of_paper = [0, 0 , img.shape[1], img.shape[0]]
    # TABLE EXTRACTION
    image_without_table, paper_list_coor = table_extraction(img, paper_list_coor)
    # TEXT EXTRACTION
    paper_list_coor, margin_of_paper = text_extraction(detect_text_model, image_without_table, paper_list_coor)
    # GET LINE OF TEXT
    paper_list_coor = get_line_of_text_from_coor(img, paper_list_coor)

    # cv2.imshow("text",image_without_table)
    # cv2.waitKey(0)

    # MAKE DOC FILE
    doc = docx.Document() 
    doc = make_doc_file(img, paper_list_coor,detect_text_model, read_text_model, doc, margin_of_paper)
    doc.save(save_dir+'/'+file_name+'.docx') 

    # for i in paper_list_coor:
    #     print(i)
    #     for j in i[1]:
    #         print(j)
    #         cv2.rectangle(image_without_line_of_text, (j[0],j[1]), (j[2],j[3]), (0, 255, 0) , 1)
    #         cv2.imshow("text",image_without_line_of_text)
    #         cv2.waitKey(0)
    # cv2.imwrite('result/line_detect.jpg',image_without_line_of_text)

 
    # cv2.imwrite('result/pre.jpg',image)
    # cv2.imwrite('result/table.jpg',img)
    # cv2.imwrite('result/without_table.jpg',image_without_table)
    # cv2.imwrite('result/line_detect.jpg',image_without_line_of_text)

def ResizeWithAspectRatio(image, max_width=None, max_height=None, min_width=None, min_height=None):
    dim = None
    (h, w) = image.shape[:2]
    print( str(h) + " " + str(w))
    image = cv2.fastNlMeansDenoisingColored(image, None, 10, 10, 7, 15) 
    if max_width is None and max_height is None and min_width is None and min_height is None:
        return image
    else:
        if max_width is None and max_height <= h:
            r = max_height / float(h)
            dim = (int(w * r), max_height)
            inter = cv2.INTER_AREA
        elif max_height is None and max_width <= w:
            r = max_width / float(w)
            dim = (max_width, int(h * r))
            inter = cv2.INTER_AREA
        else: 
            if min_width is None and min_height >= h:
                dim = (int(min_height*2), min_height)
                inter = cv2.INTER_LINEAR
            elif min_height is None and min_width >= w:
                dim = (min_width, int(min_width*2))
                inter = cv2.INTER_LINEAR
            else:
                return image

        return cv2.resize(image, dim, interpolation=inter)
    
def get_image_to_doc(detect_model, read_model, image, file_name, save_path='result'):
    # im1 = cv2.imread(file_path, 1)
    # try_rotate_image = rotate_image((ResizeWithAspectRatio(im1, height=800)))
    # try:
        # if -30 < try_rotate_image[1] < 30:
        #     from_image_to_docx(detect_model,read_model, try_rotate_image[0],file_name, save_path)
        # else:
    from_image_to_docx(detect_model,read_model, (ResizeWithAspectRatio(image, max_height=1200, min_height=800)),file_name, save_path)
    # except Exception as e:
    #     return e
    return "IMAGE TO DOC SUCCESSFULLY"
    
#-----------------Main------------------------
# if __name__ == '__main__':
#     model = easyocr.Reader(['vi'], gpu=False)
#     config = Cfg.load_config_from_name(r'vgg_seq2seq') # sử dụng config mặc định của mình 
#     config['weights'] = r'lib\vietocr\weights\vgg_seq2seq.pth' # đường dẫn đến trọng số đã huấn luyện hoặc comment để sử dụng pretrained model của mình
#     config['device'] = 'cpu' # device chạy 'cuda:0', 'cuda:1', 'cpu'
#     detector = Predictor(config)
#     file =  r'test_img/33.jpg'
#     im1 = cv2.imread(file, 1)
#     try_rotate_image = rotate_image((ResizeWithAspectRatio(im1, height=800)))
#     if -30 < try_rotate_image[1] < 30:
#         from_image_to_docx(model,detector, try_rotate_image[0], 'result')
#     else:
#         from_image_to_docx(model,detector, (ResizeWithAspectRatio(im1, height=800)), 'result')